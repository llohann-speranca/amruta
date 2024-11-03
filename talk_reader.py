from docx import Document
import os
from csv import DictWriter
from datetime import datetime
import re
from tqdm import tqdm
import traceback


# def get_metadata_file(path="./data/metadata/2024-02-xx Maha talks list .xlsx"):
#     import pandas as pd
#     metadata_df = pd.read_excel(path, index_col=0)
#     return metadata_df


def parse_metadata_from_talk(talk: list[str]):
    # [Title, Type of talk, Location, Talks status, Number of words]
    # Example talk:
    # 04 May 1982
    # Shri Mataji Teaches Yogis to Sing Bhajans in Ashram in Le Raincy
    # Talk to Sahaja Yogis on Eve of Sahasrara Puja
    # Le Raincy, France
    # Talk Language: English | Transcript: Draft
    while len(talk[0]) < 2:
        talk.pop(0)
    metadata_list = talk[:5]
    number_of_words = sum([len(paragraph.split(" ")) for paragraph in talk])
    metadata = {
        "Date": re.sub(
            "[^A-Za-z0-9 ]",
            "",
            metadata_list[0]
            .strip()
            .replace("th", "")
            .replace("  ", " ")
            .replace("1001", ""),
        ),
        "Title": metadata_list[1],
        "Type of talk": metadata_list[2],
        "Location": metadata_list[3],
        "Talk Language": metadata_list[-1]
        .split(" | ")[0]
        .removeprefix("Talk Language: "),
        "Talks status": metadata_list[-1].split(" | ")[-1].removeprefix("Transcript: "),
        "Number of words": number_of_words,
    }
    return metadata


class TalkReader:
    def __init__(
        self,
        list_of_talks_path: str,
        metadata_file_path: str,
        save_folder: str = "data/output",
    ):
        # print("Initiated")
        self.document = Document(list_of_talks_path)
        self.metadata = []
        self.metadata_file_path = metadata_file_path
        self.talks = []
        self.save_folder = save_folder
        os.makedirs(save_folder, exist_ok=True)
        self.metadata_fields = [
            "Date",
            "Title",
            "Type of talk",
            "Location",
            "Talks status",
            "Number of words",
            "Talk Language",
            "Comments",
        ]
        self.parsing_that_went_wrong = []
        self.paragraphs_do_not_starting_with_hashtag = []
        if not os.path.exists(metadata_file_path):
            os.makedirs(os.path.dirname(metadata_file_path), exist_ok=True)
            with open(self.metadata_file_path, "w") as metadata_file:
                csv_writer = DictWriter(metadata_file, fieldnames=self.metadata_fields)
                csv_writer.writeheader()

        # self.talks = self.process_document()

    def _get_all_text(self):
        document_text = [p.text for p in self.document.paragraphs]
        document_text = "\n\n".join(document_text)
        return document_text

    def _split_text(self, text: str):
        split_text = [t.strip().removeprefix("\u2003") for t in text.split("##")]
        split_text = [t for t in split_text if t]
        split_text = [t.split("\n\n") for t in split_text]
        return split_text

    def _add_talks_to_database(self, talks: list[list[str]]):
        for talk in tqdm(talks):
            self.add_talk_to_database(talk)

    def process_document(self):
        text = self._get_all_text()
        talks = self._split_text(text)
        self._add_talks_to_database(talks)
        self.dump_metadata()
        return self.metadata

    # def process_document(self):
    #     # document = [p.text for p in document.paragraphs]
    #     talk = []
    #     if os.getenv("DEBUG") == "true":
    #         paragraphs = self.document.paragraphs[:1000]
    #     else:
    #         paragraphs = self.document.paragraphs

    #     for paragraph in tqdm(paragraphs):
    #         text: str = paragraph.text
    #         if text.startswith("##"):
    #             self.add_talk_to_database(talk)
    #             text = text.removeprefix("##").removeprefix("\u2003").strip()
    #             if text:
    #                 talk = [text]
    #         # elif '##' in text:
    #         #     self.paragraphs_do_not_starting_with_hashtag.append(text)
    #         else:
    #             talk.append(text)
    #     self.add_talk_to_database(talk)
    #     return self.metadata

    def add_talk_to_database(self, talk):
        try:
            if len(talk) < 5:
                self.parsing_that_went_wrong.append(
                    {
                        "talk": talk[:10],
                        "error": "len(talk) < 5",
                        "traceback": "",
                    }
                )

            else:
                metadata = self.parse_metadata_from_talk(talk)
                self.metadata.append(metadata)
                self.dump_talk(talk=talk, metadata=metadata)
                # self.talks.append(talk)
        except Exception as e:
            self.parsing_that_went_wrong.append(
                {
                    "talk": talk[:10],
                    "error": str(e),
                    "traceback": traceback.format_exc(),
                }
            )

    def dump_talk(self, talk: list[str], metadata: dict, save_folder=None):
        if save_folder is None:
            save_folder = self.save_folder
        document = Document()
        # style = document.styles['Normal']
        # font = style.font
        # font.name = 'Calibri'
        # font.size = Pt(11)

        for paragraph in talk:
            _ = document.add_paragraph(paragraph)
            # para.style = document.styles['Normal']
        date_iso = metadata["Date"]
        save_file_name = (
            f"{date_iso} – {metadata['Title']} – {metadata['Location']}.docx"
        )
        document.save(os.path.join(save_folder, save_file_name))

    @staticmethod
    def _extract_date(date_string: str) -> tuple[str, str]:
        date_string = date_string.replace("Date Unknown", "1 January 0001")
        date_string = re.sub(
            "[^A-Za-z0-9 ]|1001|th", "", date_string.strip().replace("  ", " ")
        )
        date_string = date_string.replace("Octobre", "October")
        date_string = date_string.replace("199\n", "0000")
        date_string = date_string.replace("200\n", "0000")
        date_string = date_string.replace("Year Unknown", "0000")
        try:
            return datetime.strptime(date_string, "%d %B %Y").strftime("%Y-%m-%d"), ""
        except ValueError:
            try:
                return datetime.strptime(date_string, "%B %d %Y").strftime(
                    "%Y-%m-%d"
                ), ""
            except ValueError:
                return date_string, "Date format not recognized"

    def parse_metadata_from_talk(self, talk: list[str]):
        # [Title, Type of talk, Location, Talks status, Number of words]
        # Example talk:
        # 04 May 1982
        # Shri Mataji Teaches Yogis to Sing Bhajans in Ashram in Le Raincy
        # Talk to Sahaja Yogis on Eve of Sahasrara Puja
        # Le Raincy, France
        # Talk Language: English | Transcript: Draft
        while len(talk[0]) < 2:
            talk.pop(0)
        metadata_list = talk[:5]
        number_of_words = sum([len(paragraph.split(" ")) for paragraph in talk])
        metadata = {
            "Title": metadata_list[1],
            "Type of talk": metadata_list[2],
            "Location": metadata_list[3],
            "Talk Language": metadata_list[-1]
            .split(" | ")[0]
            .removeprefix("Talk Language: "),
            "Talks status": metadata_list[-1]
            .split(" | ")[-1]
            .removeprefix("Transcript: "),
            "Number of words": number_of_words,
        }

        date_iso, comments = self._extract_date(metadata_list[0])
        metadata["Date"] = date_iso
        metadata["Comments"] = comments
        return metadata

    def dump_metadata(self):
        with open(self.metadata_file_path, "a") as metadata_file:
            csv_writer = DictWriter(metadata_file, fieldnames=self.metadata_fields)
            csv_writer.writerows(self.metadata)


# def break_document
